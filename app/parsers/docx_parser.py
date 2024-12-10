import os
from dataclasses import dataclass
from io import BytesIO
from typing import Optional, Union

from docx import Document
from tqdm import tqdm
import subprocess
from pathlib import Path
from app.core.progress import ProgressTracker
from app.core.translator import Translator, TranslationPreferences
from app.utils.logger import setup_logger

# from docx2pdf import convert

@dataclass
class DocParserConfig:
    skip_headers: bool = True
    skip_footers: bool = True
    skip_tables: bool = False


class DocParser:
    def __init__(self,
                 translator: Translator,
                 config: Optional[DocParserConfig] = None,
                 progress_tracker: Optional[ProgressTracker] = None):
        self.translator = translator
        self.config = config or DocParserConfig()
        self.progress_tracker = progress_tracker
        self.logger = setup_logger(self.__class__.__name__)

    def _count_total_runs(self, doc: Document) -> int:
        """统计文档中需要处理的总runs数"""
        total = 0

        # 统计段落
        for para in doc.paragraphs:
            total += len([run for run in para.runs if run.text.strip()])

        if not self.config.skip_tables:
            # 统计表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            total += len([run for run in para.runs if run.text.strip()])

        if not self.config.skip_headers:
            # 统计页眉
            for section in doc.sections:
                for para in section.header.paragraphs:
                    total += len([run for run in para.runs if run.text.strip()])

        if not self.config.skip_footers:
            # 统计页脚
            for section in doc.sections:
                for para in section.footer.paragraphs:
                    total += len([run for run in para.runs if run.text.strip()])

        return total

    async def translate_document(self,
                                 doc_source: Union[str, BytesIO],
                                 filename: Optional[str] = None,
                                 output_path: Optional[str] = None,
                                 preferences: Optional[TranslationPreferences] = None) -> BytesIO:
        self.logger.info("开始翻译文档...")
        """
        在文档上直接进行翻译

        Args:
            doc_source: Word文档路径或BytesIO对象
            filename: 文件名（用于显示）
            output_path: 输出文件路径（可选）
            preferences: 翻译偏好设置

        Returns:
            BytesIO: 包含翻译后文档的BytesIO对象
        """
        # 打开文档
        doc = Document(doc_source)
        self.logger.info("文档已打开")

        # 获取总运行数并设置进度追踪器
        total_runs = self._count_total_runs(doc)
        if self.progress_tracker:
            self.progress_tracker.set_total(total_runs)

        self.logger.info(f"总运行数: {total_runs}")

        # 收集所有需要翻译的文本和位置信息
        texts_to_translate = []
        text_locations = []

        # 收集正文段落的文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                current_texts = []
                current_runs = []
                for run in paragraph.runs:
                    if run.text.strip():
                        current_texts.append(run.text)
                        current_runs.append(run)
                if current_texts:
                    texts_to_translate.append(" ".join(current_texts))
                    text_locations.append(('paragraph', current_runs))

        # 收集表格中的文本
        if not self.config.skip_tables:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                current_texts = []
                                current_runs = []
                                for run in paragraph.runs:
                                    if run.text.strip():
                                        current_texts.append(run.text)
                                        current_runs.append(run)
                                if current_texts:
                                    texts_to_translate.append(" ".join(current_texts))
                                    text_locations.append(('table', current_runs))

        # 收集页眉文本
        if not self.config.skip_headers:
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        current_texts = []
                        current_runs = []
                        for run in paragraph.runs:
                            if run.text.strip():
                                current_texts.append(run.text)
                                current_runs.append(run)
                        if current_texts:
                            texts_to_translate.append(" ".join(current_texts))
                            text_locations.append(('header', current_runs))

        # 收集页脚文本
        if not self.config.skip_footers:
            for section in doc.sections:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        current_texts = []
                        current_runs = []
                        for run in paragraph.runs:
                            if run.text.strip():
                                current_texts.append(run.text)
                                current_runs.append(run)
                        if current_texts:
                            texts_to_translate.append(" ".join(current_texts))
                            text_locations.append(('footer', current_runs))

        self.logger.info("文本收集完成，准备翻译...")

        # 使用进度条显示翻译进度
        with tqdm(total=total_runs, desc=f"正在翻译: {filename}") as pbar:
            # 批量翻译
            chunk_size = 10  # 每次翻译10个文本块
            for i in range(0, len(texts_to_translate), chunk_size):
                chunk_texts = texts_to_translate[i:i + chunk_size]
                chunk_locations = text_locations[i:i + chunk_size]

                # 批量翻译
                translations = await self.translator.batch_translate(chunk_texts, preferences)

                # 更新文档
                for text, location in zip(translations, chunk_locations):
                    loc_type, runs = location
                    # 计算本次更新涉及的run 数量
                    runs_count = len(runs)

                    # 不再使用分词的方式，而是直接替换整个文本
                    if len(runs) == 1:
                        # 如果只有一个run，直接替换
                        runs[0].text = text
                    else:
                        # 如果有多个run，尝试按比例分配翻译后的文本
                        total_original_length = sum(len(run.text) for run in runs)
                        if total_original_length == 0:
                            continue
                            
                        # 将翻译后的文本按原文本的长度比例分配给每个run
                        current_pos = 0
                        for run in runs:
                            original_length = len(run.text)
                            if original_length == 0:
                                continue
                                
                            # 计算这个run应该获得的翻译文本的比例
                            proportion = original_length / total_original_length
                            text_length = int(len(text) * proportion)
                            
                            # 确保最后一个run获得所有剩余文本
                            if run is runs[-1]:
                                run.text = text[current_pos:]
                            else:
                                run.text = text[current_pos:current_pos + text_length]
                                current_pos += text_length

                    # 更新进度条和进度追踪器
                    pbar.update(runs_count)
                    if self.progress_tracker:
                        self.progress_tracker.update(runs_count)

        # 保存docx文档
        output_buffer = BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)

        # 如果指定了输出路径，同时保存到文件
        if output_path:
            doc.save(output_path)
            pdf_path = output_path.replace('.docx', '.pdf')
            try:
                subprocess.run([
                    '/usr/local/bin/soffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', str(Path(output_path).parent),
                    output_path
                ], check=True)
                self.logger.info(f"PDF转换成功: {pdf_path}")
            except subprocess.CalledProcessError as e:
                self.logger.error(f"PDF转换失败: {e}")

        return output_buffer